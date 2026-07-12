"""
Generates the MVC-ARENA study questionnaire as a .docx file.

Run:  python questionnaire.py   (from the supportdocs folder)
Output: MVC-ARENA_Questionnaire.docx
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, level=1):
    return doc.add_heading(text, level=level)


def para(text="", bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def mcq(num, stem, options):
    """A multiple-choice question: numbered stem + a,b,c,d options."""
    p = doc.add_paragraph()
    p.add_run(f"{num}. ").bold = True
    p.add_run(stem)
    letters = "abcd"
    for i, opt in enumerate(options):
        op = doc.add_paragraph()
        op.paragraph_format.left_indent = Pt(24)
        op.add_run(f"{letters[i]}) ").bold = True
        op.add_run(opt)


def likert(num, stem):
    p = doc.add_paragraph()
    p.add_run(f"{num}. ").bold = True
    p.add_run(stem)
    s = doc.add_paragraph("   1   2   3   4   5")
    s.paragraph_format.left_indent = Pt(24)


# ================= TITLE =================
title = doc.add_heading("MVC-ARENA Study Questionnaire", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

para(
    "This questionnaire has two parts: a qualitative survey (demographics and "
    "perceived value) and a quantitative MVC knowledge test. The knowledge test "
    "comes in three parallel sets (A, B, C) that share one blueprint, so question "
    "N tests the same category in all three sets.",
    italic=True,
)

heading("How to administer", level=2)
bullet("Quantitative test: give one set before playing (pre-test) and a different "
       "set after (post-test); optionally a third set later (retention). Rotate "
       "which set is used as pre/post across participants to remove order effects.")
bullet("Scoring: 1 point per item (max 20). Learning gain = post-test minus pre-test.")
bullet("Qualitative survey: administer once, after playing.")

# ================= PART 1: QUALITATIVE =================
heading("Part 1 - Qualitative Survey", level=1)

heading("A. Demographics & background", level=2)
mcq(1, "Age range:", ["under 18", "18-24", "25-34", "35 or older"])
mcq(2, "Gender (optional):", ["Female", "Male", "Prefer not to say", "Other: ______"])
mcq(3, "Education Level:", ["High School", "Bachelors", "Masters", "Other: ______"])
mcq(4, "Programming experience:", ["None", "Less than 1 year", "1-3 years", "3+ years"])
mcq(5, "Had you heard of the MVC (Model-View-Controller) concept before this study?",
    ["Never", "Heard of it but did not understand it",
     "Understood it somewhat", "Understood it well"])
mcq(6, "Familiarity with object-oriented programming (inheritance, interface, polymorphism):",
    ["None", "Basic", "Comfortable", "Advanced"])
mcq(7, "How often do you play card or video games?",
    ["Never", "Rarely", "Sometimes", "Often"])

heading("B. Perceived value & experience", level=2)
para("Rate each statement from 1 (Strongly disagree) to 5 (Strongly agree).", italic=True)
likert(8, "The game was easy to learn and play.")
likert(9, "The game was engaging and enjoyable.")
likert(10, "The game helped me understand what MVC is.")
likert(11, "The game helped me understand which components belong to Model, View, and Controller.")
likert(12, "The game helped me understand security threats and their defenses.")
likert(13, "The game helped me understand OOP ideas (inheritance, interface, polymorphism).")
likert(14, "The video tutorial was helpful more?")
likert(15, "The instruction manual was helpful more?")
likert(16, "The difficulty level was appropriate.")
likert(17, "I feel the game was a valuable way to learn these concepts.")
likert(18, "I would recommend this game to others learning MVC.")
likert(19, "The game is time-consuming.")
likert(20, "Playing against bots is better than playing with humans.")

heading("C. Open-ended", level=2)
para("21. What did you find most valuable about the game for learning?")
para("____________________________________________________________________")
para("22. What was confusing or unclear?")
para("____________________________________________________________________")
para("23. What would you change or add?")
para("____________________________________________________________________")

# ================= PART 2: QUANTITATIVE =================
doc.add_page_break()
heading("Part 2 - Quantitative MVC Knowledge Test", level=1)

heading("Blueprint (category by question number - same for Sets A, B, C)", level=2)
blueprint = [
    ("1", "What MVC is (definition / purpose)"),
    ("2", "Role of the Model"),
    ("3", "Role of the View"),
    ("4", "Role of the Controller"),
    ("5", "Classify a component -> Model"),
    ("6", "Classify a component -> View"),
    ("7", "Classify a component -> Controller"),
    ("8", "Benefit of MVC (separation of concerns)"),
    ("9", "Inheritance (OOP)"),
    ("10", "Polymorphism (OOP)"),
    ("11", "Security defense mapping (attack -> defense)"),
    ("12", "Security defense mapping (attack -> defense)"),
    ("13", "Applied scenario (component + layer)"),
    ("14", "MVC request flow (correct order)"),
    ("15", "Flow: what the Controller does after receiving a request"),
    ("16", "Flow: the View shows the final result to the user"),
    ("17", "Git / version control"),
    ("18", "Error handling"),
    ("19", "Logging (Logger)"),
    ("20", "Disaster & recovery"),
]
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
tbl.rows[0].cells[0].paragraphs[0].add_run("Q").bold = True
tbl.rows[0].cells[1].paragraphs[0].add_run("Category").bold = True
for q, cat in blueprint:
    row = tbl.add_row().cells
    row[0].text = q
    row[1].text = cat

# ---- the three sets (category order matches the blueprint above) ----
SETS = {
    "SET A": [
        ("What does MVC stand for?",
         ["Model-View-Controller", "Module-Variable-Class",
          "Method-View-Component", "Model-Variable-Controller"]),
        ("Which layer is mainly responsible for the application's data and business logic?",
         ["View", "Controller", "Model", "Router"]),
        ("Which layer is responsible for presenting information to the user?",
         ["Model", "View", "Controller", "Database"]),
        ("Which layer receives user requests and coordinates the Model and View?",
         ["View", "Model", "Controller", "Interface"]),
        ("A Database component belongs to which layer?",
         ["Model", "View", "Controller", "None"]),
        ("A Web View component belongs to which layer?",
         ["Model", "View", "Controller", "Network"]),
        ("An Authentication component belongs to which layer?",
         ["Model", "View", "Controller", "Database"]),
        ("The main advantage of separating an app into Model, View, and Controller is that it:",
         ["makes the program run faster",
          "separates concerns so each part can change independently",
          "removes the need for a database", "guarantees security"]),
        ("Inheritance in OOP means:",
         ["one class reuses and extends another's behavior", "two classes run at once",
          "hiding a class's data", "splitting a program into files"]),
        ("Polymorphism allows:",
         ["one name/interface to take many forms", "only one object at a time",
          "code that never changes", "data to be encrypted"]),
        ("Which defense protects against SQL Injection?",
         ["ORM", "Web View", "Logger", "Rate Limiting"]),
        ("Which defense protects against Unauthorized Access?",
         ["Caching", "Authentication/Authorization", "ORM", "Output Validation"]),
        ("You need to safely store passwords and API keys. Best fit?",
         ["Secrets Manager - Model", "Web View - View", "Routing - Controller", "Logger - View"]),
        ("Put the MVC request cycle in the correct order:",
         ["View -> Model -> Controller",
          "Controller -> Model -> Controller -> View",
          "Model -> View -> Controller",
          "View -> Controller -> Model"]),
        ("When the Controller receives a user request, what does it typically do next?",
         ["Render the final page itself", "Ask the Model for the needed data",
          "Send the request back to the browser", "Delete the database"]),
        ("In the request cycle, the View's job at the end is to:",
         ["store the data in the database",
          "take the data the Controller provides and present it to the user",
          "decide the application's business rules", "route the next request"]),
        ("Your team's latest changes broke the app. Which practice lets you go back "
         "to a working earlier version?",
         ["Git version control", "Caching", "Rate limiting", "an ORM"]),
        ("A program receives unexpected input. With good error handling, the program will:",
         ["crash immediately", "catch the problem and continue or fail safely",
          "delete its data", "ignore the user"]),
        ("In the game a Logger defends against a Bug. In real software, logging mainly helps you:",
         ["find and diagnose defects", "encrypt data",
          "render the user interface", "limit requests"]),
        ("In the game, a Disaster represents a catastrophic failure or data loss. "
         "Which practice best helps you recover?",
         ["Git backups", "Web View", "Rate Limiting", "Caching"]),
    ],
    "SET B": [
        ("MVC is mainly a software design pattern used to:",
         ["encrypt traffic",
          "organize an application into three connected parts with separate responsibilities",
          "compress images", "schedule CPU tasks"]),
        ("Business rules and application data live mainly in the:",
         ["Controller", "Model", "View", "Middleware"]),
        ("The part the user actually sees and interacts with is the:",
         ["Model", "Controller", "View", "ORM"]),
        ("Handling an incoming request and deciding what to do is the job of the:",
         ["Model", "Controller", "View", "Cache"]),
        ("An ORM component belongs to which layer?",
         ["View", "Controller", "Model", "None"]),
        ("A Mobile View component belongs to which layer?",
         ["Model", "Controller", "View", "Database"]),
        ("A Routing component belongs to which layer?",
         ["Model", "View", "Controller", "Storage"]),
        ("Separation of concerns in MVC mainly helps to:",
         ["make code easier to maintain and change", "reduce files to exactly three",
          "eliminate all bugs", "make the UI colorful"]),
        ("A class 'Dog' that reuses and extends class 'Animal' is an example of:",
         ["Polymorphism", "Inheritance", "Encapsulation", "Routing"]),
        ("Calling the same method name on different objects, each behaving appropriately, is:",
         ["Inheritance", "Polymorphism", "Caching", "Logging"]),
        ("Which defense protects against CSRF?",
         ["CSRF Protection", "ORM", "Caching", "Logger"]),
        ("Which defense protects against a Bug hazard?",
         ["Git", "Logger", "ORM", "Web View"]),
        ("You want to limit how many requests a client can send to prevent overload. Best fit?",
         ["Rate Limiting - Controller", "Database - Model", "CLI View - View", "ORM - Model"]),
        ("Which sequence best describes how a request flows in MVC?",
         ["The request goes to the Controller, which uses the Model, then passes results to the View",
          "The request goes straight to the database and back",
          "The View calls the Model, which calls the Controller",
          "The Model receives the request first, then the View"]),
        ("Immediately after receiving a request, the Controller usually:",
         ["updates the user interface directly", "works with the Model to get or change data",
          "shuts down the server", "encrypts the View"]),
        ("Once the Controller has the data, it chooses a View to:",
         ["permanently store the data", "display the result to the user",
          "authenticate the user", "route the request again"]),
        ("A teammate asks 'what changed in the code last week, and who changed it?' "
         "Which tool answers this?",
         ["Git version control", "an ORM", "a Web View", "Rate Limiting"]),
        ("In the game, Error Handling is a defensive card. In real software it mainly improves:",
         ["the program's stability and robustness when something goes wrong",
          "the screen colors", "the network speed only", "the number of files"]),
        ("A hard-to-reproduce bug appears. What best helps you trace what happened?",
         ["logs from a Logger", "a faster database", "a new View", "rate limiting"]),
        ("If a Disaster (data loss) happens, what lets you restore your work?",
         ["a Logger", "Git version control / backups", "an ORM", "Output Validation"]),
    ],
    "SET C": [
        ("Which best describes the MVC pattern?",
         ["splitting an application into Model, View, and Controller with distinct roles",
          "a programming language", "a type of database", "a security certificate"]),
        ("Where would you put logic that validates and saves a user record to storage?",
         ["View", "Controller", "Model", "Router"]),
        ("A command-line interface (CLI) that shows output to the user is part of the:",
         ["Model", "View", "Controller", "Cache"]),
        ("Middleware that checks each request before it reaches the main logic belongs to the:",
         ["Model", "View", "Controller", "Database"]),
        ("A Caching component belongs to which layer?",
         ["Model", "View", "Controller", "None"]),
        ("An Output Validation (escaping data before display) component belongs to which layer?",
         ["Model", "View", "Controller", "Storage"]),
        ("An Authorization component belongs to which layer?",
         ["Model", "View", "Controller", "Database"]),
        ("A key benefit of MVC is that you can change the View without rewriting the:",
         ["business logic in the Model", "operating system",
          "programming language", "network cables"]),
        ("Inheritance helps mainly by:",
         ["letting a subclass reuse common code from a parent class",
          "encrypting the parent", "running classes in parallel", "hiding the program"]),
        ("Polymorphism is useful because it lets you:",
         ["treat different types through one common interface", "store more data",
          "prevent inheritance", "speed up only the database"]),
        ("Which defense protects against Cross-Site Scripting (XSS)?",
         ["Data Validation", "Rate Limiting", "Git", "Authentication"]),
        ("Which defense protects against Malware?",
         ["File Storage Adapter", "Web View", "Routing", "CSRF Protection"]),
        ("A Disaster (catastrophic data loss) strikes. Best recovery practice?",
         ["Git - version control/backups", "Logger - Model", "Caching - View", "ORM - Controller"]),
        ("A user clicks a button. In MVC, the correct order of handling is:",
         ["Controller receives it -> Model provides data -> Controller passes it to the View -> View shows the user",
          "View shows the user -> Model -> Controller",
          "Model -> Controller -> Model -> View",
          "Database -> View -> Controller"]),
        ("After the Controller gets the request, the Model is responsible for:",
         ["drawing the screen", "providing or updating the data / business logic",
          "routing to another controller", "styling the page"]),
        ("The last step, showing the prepared data to the user, is done by the:",
         ["Model", "View", "Router", "Database"]),
        ("A benefit of using Git version control is that you can:",
         ["recover an earlier version of your work if something goes wrong",
          "make the app run faster", "remove the need for a Model", "prevent all attacks"]),
        ("Which is the best example of error handling?",
         ["catching an exception and showing a helpful message instead of crashing",
          "storing passwords", "routing a request", "doubling the score"]),
        ("In the game a Logger protects against a Bug. Why does that match real development?",
         ["logs help you find and fix defects", "logs encrypt the database",
          "logs draw the user interface", "logs speed up the network"]),
        ("A Disaster is best described as:",
         ["a minor typo", "a catastrophic failure or loss of data",
          "a slow network", "a user login"]),
    ],
}

for set_name, questions in SETS.items():
    doc.add_page_break()
    heading(set_name, level=2)
    para("Name/ID: ______________________    Date: ____________    "
         "(circle: pre-test / post-test / retention)", italic=True)
    for i, (stem, options) in enumerate(questions, start=1):
        mcq(i, stem, options)

# ---- answer keys ----
doc.add_page_break()
heading("Answer Keys", level=2)
keys = {
    "Set A": "1-a, 2-c, 3-b, 4-c, 5-a, 6-b, 7-c, 8-b, 9-a, 10-a, 11-a, 12-b, 13-a, 14-b, 15-b, 16-b, 17-a, 18-b, 19-a, 20-a",
    "Set B": "1-b, 2-b, 3-c, 4-b, 5-c, 6-c, 7-c, 8-a, 9-b, 10-b, 11-a, 12-b, 13-a, 14-a, 15-b, 16-b, 17-a, 18-a, 19-a, 20-b",
    "Set C": "1-a, 2-c, 3-b, 4-c, 5-a, 6-b, 7-c, 8-a, 9-a, 10-a, 11-a, 12-a, 13-a, 14-a, 15-b, 16-b, 17-a, 18-a, 19-a, 20-b",
}
for name, key in keys.items():
    p = doc.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(key)

note = doc.add_paragraph()
note.add_run("Note: ").bold = True
note.add_run(
    "Component-to-layer answers follow the game's own classification (e.g. the "
    "game places Authentication, Authorization, and Rate Limiting in the "
    "Controller lane, Output Validation in the View lane, and Caching, ORM, and "
    "Secrets Manager in the Model lane). This is appropriate for measuring what "
    "the game teaches; note in your methodology that a few of these can sit in "
    "other layers in different real-world designs."
)

out = os.path.join(os.path.dirname(__file__), "MVC-ARENA_Questionnaire.docx")
doc.save(out)
print("wrote", out)
